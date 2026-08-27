import re
import os
import sys

# === 0. 路径修正：确保能找到同级模块 ===
# 防止在不同目录下运行时找不到 latex_utils
if os.path.dirname(__file__) not in sys.path:
    sys.path.append(os.path.dirname(__file__))

# === 1. 导入 python-docx 及其组件 ===
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# === 2. 从同级模块导入 LaTeX 处理工具 ===
try:
    from latex_utils import clean_inline_latex_to_unicode, latex_to_image_stream
except ImportError:
    # 兜底：如果找不到模块，定义简单的清理逻辑
    def clean_inline_latex_to_unicode(t): return t
    def latex_to_image_stream(s): return None

# =========================================================================
# 🛠️ 核心工具函数：生成 Word (含公式图片)
# =========================================================================
def save_text_to_docx(text_content, file_path):
    """
    将 Markdown 内容保存为 Docx。
    自动将 LaTeX 公式块转换为图片插入，并严格锁定中西文字体。
    """
    if not HAS_DOCX:
        print("⚠️ Warning: python-docx not installed.")
        return False

    try:
        doc = Document()
        
        # 1. 设置全局默认样式 (Normal)
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        # 强制锁定东亚（中文）字符集使用 Times New Roman
        style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

        # 2. 逐行处理 Markdown
        lines = text_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # === 🔥 FIX 1: 强力拦截单独出现的 LaTeX 闭合符 ===
            # 先切除行末可能粘连的闭合符
            if line.endswith(r'\]'): 
                line = line[:-2].strip()
            if line.endswith(r'$$'): 
                line = line[:-2].strip()
            
            # 如果切完没内容了，或者是单独的闭合符行，直接跳过
            # 防止图片下方出现悬空的 "]" 或 "bb"
            if not line or line in [r']', r')', r'}', r'[', r'bb']: 
                continue


            # === A0. Markdown image handling: ![alt](path) ===
            img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
            if img_match:
                alt_text = img_match.group(1).strip()
                img_path = img_match.group(2).strip().strip('"').strip("'")
                if not os.path.isabs(img_path):
                    img_path = os.path.join(os.path.dirname(file_path), img_path)
                if os.path.exists(img_path):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    try:
                        run.add_picture(img_path, width=Inches(5.8))
                    except Exception:
                        run.add_picture(img_path)
                    if alt_text:
                        cap = doc.add_paragraph(alt_text)
                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in cap.runs:
                            r.font.name = 'Times New Roman'
                            r.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                            r.font.size = Pt(10)
                            r.font.italic = True
                else:
                    p = doc.add_paragraph(f"[Missing image: {img_path}]")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # === A. 标题处理 (#, ##, ###) ===
            if line.startswith('#'):
                if line.startswith('### '):
                    level, size, content = 3, 14, line.replace('### ', '')
                elif line.startswith('## '):
                    level, size, content = 2, 15, line.replace('## ', '')
                elif line.startswith('# '):
                    level, size, content = 1, 16, line.replace('# ', '')
                else:
                    level, size, content = 1, 16, line.lstrip('#').strip()

                # 标题清洗逻辑
                content = clean_inline_latex_to_unicode(content)
                heading = doc.add_heading(content, level=level)
                
                # 格式化标题样式
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    run.font.size = Pt(size)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)

            # === B. 公式块处理 (图片化) ===
            # 识别 LaTeX 块标记或数学关键字
            elif (line.startswith('\\[') or line.startswith('$$') or 
                  (line.startswith('\\') and any(k in line for k in ['text', 'frac', 'sum', 'arrow', 'cdot']))):
                
                # 调用同级 latex_utils 模块渲染
                img_stream = latex_to_image_stream(line)
                
                if img_stream:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    
                    # === 核心修改：移除强制宽度，让 Word 根据 DPI 自动缩放 ===
                    # 只要 latex_utils.py 使用了 bbox_inches='tight' 和高 DPI，
                    # 这里直接插入效果是最好的，既清晰又不会留白。
                    run.add_picture(img_stream) 
                else:
                    # === 🔥 FIX 2: 渲染失败时的深度回退清洗 ===
                    # 1. 移除公式外壳
                    clean_content = line.replace(r'\[', '').replace(r'\]', '').replace(r'$$', '')
                    # 2. 剥离 \text{...}
                    clean_content = re.sub(r'\\text\{([^}]+)\}', r'\1', clean_content)
                    # 3. 替换复杂箭头为 Unicode 箭头
                    clean_content = clean_content.replace(r'\xrightarrow', '→').replace(r'\rightarrow', '→')
                    
                    clean_content = clean_content.strip()
                    if clean_content:
                        p = doc.add_paragraph(clean_content)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            run.font.italic = True
                            run.font.name = 'Times New Roman'

            # === C. 标准正文处理 ===
            else:
                # 处理行内公式符号和 Unicode 映射
                clean_line = clean_inline_latex_to_unicode(line)
                # 二次清洗可能残留的行内 LaTeX 标记
                clean_line = re.sub(r'\\\(|\\\)|\\\[|\\\]', '', clean_line)
                
                if clean_line.strip():
                    p = doc.add_paragraph(clean_line)
                    p.paragraph_format.line_spacing = 1.5
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                        run.font.size = Pt(12)
                        run.font.color.rgb = RGBColor(0, 0, 0)

        doc.save(file_path)
        return True
    except Exception as e:
        print(f"❌ Word generation failed: {e}")
        return False