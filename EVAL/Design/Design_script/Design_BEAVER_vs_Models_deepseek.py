# -*- coding: utf-8 -*-
"""
Design-mode LLM-as-a-Judge scorer: 3 candidates per family, Grok excluded.

What this version does
----------------------
1) Scores every candidate separately:
   - BEAVER_Paper_1/2/3 from Paper_*.docx or BEAVER*.docx
   - GPT5.5-1/2/3 from gpt5.5-1/2/3.txt or similar
   - Gemini3.1-pro-1/2/3 from gemini3.1pro-1/2/3.txt or similar
   - DeepSeek-V4-pro-1/2/3 from Deepseek-V4-pro-1/2/3.txt or similar
   - Qwen3.5-Plus-1/2/3 from Qwen3.5-Plus-1/2/3.txt or similar
2) Ignores Grok files completely.
3) Runs forward/reverse joint scoring to reduce position bias.
4) Computes:
   - per-candidate averaged scores
   - per-family mean scores: BEAVER_Mean, GPT5.5_Mean, Gemini3.1-pro_Mean, DeepSeek-V4-pro_Mean, Qwen3.5-Plus_Mean
   - per-family best scores: BEAVER_Best, GPT5.5_Best, Gemini3.1-pro_Best, DeepSeek-V4-pro_Best, Qwen3.5-Plus_Best
5) Uses a seven-dimensional rubric:
   - Chemical_Feasibility
   - Scientific_Rigor
   - Mechanistic_Coherence
   - Property_Constraint_and_Tradeoff_Handling
   - Design_Completeness
   - Risk_and_Uncertainty_Awareness
   - Design_Innovation_and_Exploration

Recommended file layout
-----------------------
ROOT_DIR/
  热学+降解/
    Paper_1_xxx.docx
    Paper_2_xxx.docx
    Paper_3_xxx.docx
    gpt5.5-1.txt
    gpt5.5-2.txt
    gpt5.5-3.txt
    gemini3.1pro-1.txt
    gemini3.1pro-2.txt
    gemini3.1pro-3.txt
    Deepseek-V4-pro-1.txt
    Deepseek-V4-pro-2.txt
    Deepseek-V4-pro-3.txt
    Qwen3.5-Plus-1.txt
    Qwen3.5-Plus-2.txt
    Qwen3.5-Plus-3.txt
  力学+降解/
    ... same pattern ...

Notes
-----
- This script intentionally does NOT hard-code your API key. Set JUDGE_API_KEY as an environment variable.
- If FORCE_REJUDGE = True and output CSVs already exist, the script may reuse them. Set FORCE_REJUDGE=True to rerun.
"""

import os
import re
import json
import time
import math
import traceback
import concurrent.futures
from pathlib import Path
from typing import Dict, List, Tuple, Any

import pandas as pd
import numpy as np
from tqdm import tqdm
from openai import OpenAI

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    HAS_MPL = True
except Exception:
    HAS_MPL = False

try:
    import seaborn as sns
    HAS_SEABORN = True
except Exception:
    HAS_SEABORN = False

try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False


# =========================================================================
# 1. User configuration
# =========================================================================

ROOT_DIR = Path(__file__).resolve().parent.parent / "Judge_Input"

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "Judge_output" / "_DesignJudge_V4-1"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

RUN_LOG_PATH = OUTPUT_DIR / "0_run_log.txt"
DISCOVERED_CASES_PATH = OUTPUT_DIR / "0_discovered_cases.csv"
RAW_OUTPUTS_PATH = OUTPUT_DIR / "1_raw_design_outputs.csv"
FORWARD_SCORE_PATH = OUTPUT_DIR / "2_item_scores_forward.csv"
REVERSE_SCORE_PATH = OUTPUT_DIR / "2_item_scores_reverse.csv"
AVERAGE_SCORE_PATH = OUTPUT_DIR / "3_item_scores_average_per_candidate.csv"
FINAL_BY_CANDIDATE_PATH = OUTPUT_DIR / "4_final_scores_by_candidate.csv"
FINAL_BY_FAMILY_MEAN_PATH = OUTPUT_DIR / "5_final_scores_by_family_mean.csv"
FINAL_BY_FAMILY_BEST_PATH = OUTPUT_DIR / "6_final_scores_by_family_best.csv"
OVERALL_MATRIX_CANDIDATE_PATH = OUTPUT_DIR / "7_case_candidate_overall_matrix.csv"
OVERALL_MATRIX_FAMILY_MEAN_PATH = OUTPUT_DIR / "8_case_family_mean_overall_matrix.csv"
HEATMAP_CANDIDATE_PATH = OUTPUT_DIR / "9_metric_heatmap_by_candidate.png"
HEATMAP_FAMILY_MEAN_PATH = OUTPUT_DIR / "10_metric_heatmap_by_family_mean.png"
OVERALL_CANDIDATE_PLOT_PATH = OUTPUT_DIR / "11_overall_bar_by_candidate.png"
OVERALL_FAMILY_MEAN_PLOT_PATH = OUTPUT_DIR / "12_overall_bar_by_family_mean.png"
CRITIQUE_PATH = OUTPUT_DIR / "13_judge_critiques.csv"
EXPLORATION_PATH = OUTPUT_DIR / "14_family_similarity_and_exploration.csv"
FAMILY_MEAN_WITH_EXPLORATION_PATH = OUTPUT_DIR / "15_final_scores_by_family_mean_with_exploration.csv"

DEFAULT_BASE_URL = os.getenv("LLM_BASE_URL", "https://llmapi.paratera.com/v1")

def env_or_default(key: str, default: str = "") -> str:
    val = os.getenv(key, default)
    return val.strip() if isinstance(val, str) else val

JUDGE_CONFIG = {
    "model_name": env_or_default("JUDGE_MODEL", "DeepSeek-V4-Pro"),
    "api_key": env_or_default("JUDGE_API_KEY"),
    "base_url": env_or_default("JUDGE_BASE_URL", DEFAULT_BASE_URL),
}

MAX_WORKERS_JUDGE = 2
SAVE_EVERY_N_JUDGE = 2
MAX_CHARS_PER_CANDIDATE = 10000
FORCE_REJUDGE = True

FAMILY_ORDER = [
    "BEAVER",
    "GPT5.5",
    "Gemini3.1-pro",
    "DeepSeek-V4-pro",
    "Qwen3.5-Plus",
]

CANDIDATE_ORDER_FORWARD = [
    "BEAVER_Paper_1",
    "BEAVER_Paper_2",
    "BEAVER_Paper_3",
    "GPT5.5-1",
    "GPT5.5-2",
    "GPT5.5-3",
    "Gemini3.1-pro-1",
    "Gemini3.1-pro-2",
    "Gemini3.1-pro-3",
    "DeepSeek-V4-pro-1",
    "DeepSeek-V4-pro-2",
    "DeepSeek-V4-pro-3",
    "Qwen3.5-Plus-1",
    "Qwen3.5-Plus-2",
    "Qwen3.5-Plus-3",
]
CANDIDATE_ORDER_REVERSE = list(reversed(CANDIDATE_ORDER_FORWARD))

FAMILY_MEAN_ORDER = [f"{f}_Mean" for f in FAMILY_ORDER]
FAMILY_BEST_ORDER = [f"{f}_Best" for f in FAMILY_ORDER]

METRIC_COLUMNS = [
    "Chemical_Feasibility",
    "Scientific_Rigor",
    "Mechanistic_Coherence",
    "Property_Constraint_and_Tradeoff_Handling",
    "Design_Completeness",
    "Risk_and_Uncertainty_Awareness",
    "Design_Innovation_and_Exploration",
]

WEIGHTS = {
    "Chemical_Feasibility": 0.17,
    "Scientific_Rigor": 0.15,
    "Mechanistic_Coherence": 0.15,
    "Property_Constraint_and_Tradeoff_Handling": 0.20,
    "Design_Completeness": 0.13,
    "Risk_and_Uncertainty_Awareness": 0.10,
    "Design_Innovation_and_Exploration": 0.10,
}

ITEM_SCORE_COLUMNS = [
    "case_id",
    "case_name",
    "case_dir",
    "design_request",
    "OrderRun",
    "Candidate",
    "Family",
    *METRIC_COLUMNS,
    "Overall",
    "JudgeParseOK",
    "Critique",
]

CASE_REQUEST_OVERRIDES: Dict[str, str] = {
    "热学+降解": (
        "Design a biodegradable polyester with improved heat resistance but without excessively suppressing degradation kinetics."
    ),
    "力学+降解": (
        "Design a biodegradable implant polymer that provides high initial mechanical support while retaining strength during early-stage hydrolytic degradation."
    ),
    "力学+热学": (
        "Design a PLA-based copolymer that increases Tg while maintaining ductility for biodegradable packaging."
    ),
    "力学+热学+降解": (
        "Design a biodegradable copolymer that simultaneously increases Tg, preserves toughness, and maintains a clinically acceptable degradation rate."
    ),
}

CASE_REQUEST_OVERRIDES.update({
    "Therm._Deg": CASE_REQUEST_OVERRIDES["热学+降解"],
    "Mech._Deg": CASE_REQUEST_OVERRIDES["力学+降解"],
    "Mech._Therm": CASE_REQUEST_OVERRIDES["力学+热学"],
    "Mech._Therm._Deg": CASE_REQUEST_OVERRIDES["力学+热学+降解"],
})

REQUEST_KEYWORDS = [
    (
        ["力学+热学+降解", "mechanical+thermal+degradation", "simultaneously", "toughness", "clinically acceptable"],
        CASE_REQUEST_OVERRIDES["力学+热学+降解"],
    ),
    (
        ["力学+热学", "PLA", "pla", "packaging", "ductility", "Tg while maintaining ductility"],
        CASE_REQUEST_OVERRIDES["力学+热学"],
    ),
    (
        ["力学+降解", "mechanical+degradation", "initial mechanical support", "early-stage hydrolytic"],
        CASE_REQUEST_OVERRIDES["力学+降解"],
    ),
    (
        ["热学+降解", "thermal+degradation", "heat", "degradation", "improved_heat_resistance"],
        CASE_REQUEST_OVERRIDES["热学+降解"],
    ),
]

IGNORED_OUTPUT_KEYWORDS = ["grok"]


# =========================================================================
# 2. Logging and utilities
# =========================================================================


def now_str() -> str:
    return time.strftime("%Y-%m-%d %H:%M:%S")


def log_msg(msg: str, also_print: bool = True) -> None:
    line = f"[{now_str()}] {msg}"
    with open(RUN_LOG_PATH, "a", encoding="utf-8") as f:
        f.write(line + "\n")
    if also_print:
        print(line)


def save_df_atomic(df: pd.DataFrame, path: Path) -> None:
    path = Path(path)
    tmp_path = path.with_suffix(path.suffix + ".tmp")
    df.to_csv(tmp_path, index=False, encoding="utf-8-sig")
    os.replace(tmp_path, path)


def natural_key(s: str):
    return [int(t) if t.isdigit() else t.lower() for t in re.split(r"(\d+)", str(s))]


def clean_think_tag(text: str) -> str:
    if text is None:
        return ""
    if not isinstance(text, str):
        text = str(text)
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()


def normalize_answer(text: str) -> str:
    text = clean_think_tag(text)
    text = text.replace("\u200b", " ").strip()
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def preview_text(text: str, max_len: int = 160) -> str:
    s = str(text).replace("\n", " ")
    return s[:max_len] + ("..." if len(s) > max_len else "")


def safe_float(x, default: float = 0.0) -> float:
    try:
        if x is None:
            return default
        val = float(x)
        if math.isnan(val):
            return default
        return val
    except Exception:
        return default


def truncate_middle(text: str, max_chars: int = MAX_CHARS_PER_CANDIDATE) -> str:
    text = "" if text is None else str(text)
    if len(text) <= max_chars:
        return text
    head = int(max_chars * 0.65)
    tail = max_chars - head
    return (
        text[:head]
        + f"\n\n[TRUNCATED MIDDLE: original length={len(text)} chars; preserved first {head} and last {tail} chars]\n\n"
        + text[-tail:]
    )


def is_invalid_answer(text: str) -> bool:
    s = "" if text is None else str(text).strip()
    if not s or s.lower() == "nan":
        return True
    s_low = s.lower()
    invalid_prefixes = [
        "error:",
        "agent error:",
        "api error:",
        "judge error:",
        "traceback (most recent call last):",
        "❌ api error:",
        "❌ error:",
        "❌ agent error:",
    ]
    if any(s_low.startswith(p) for p in invalid_prefixes):
        return True
    if len(s) < 20:
        return True
    return False


def extract_index_from_filename(path: Path) -> int:
    stem = path.stem

    # Prefer explicit suffix index such as gpt5.5-2, Paper_3, candidate_1.
    patterns = [
        r"(?:paper|candidate|idea|output|model)[_\-\s]*(\d+)",
        r"[_\-\s](\d+)$",
    ]
    for pat in patterns:
        m = re.search(pat, stem, flags=re.IGNORECASE)
        if m:
            return int(m.group(1))

    nums = re.findall(r"(\d+)", stem)
    if nums:
        return int(nums[-1])
    return 1


def get_family(candidate_name: str) -> str:
    s = str(candidate_name)
    if s.startswith("BEAVER"):
        return "BEAVER"
    if s.startswith("GPT5.5"):
        return "GPT5.5"
    if s.startswith("Gemini3.1-pro"):
        return "Gemini3.1-pro"
    if s.startswith("DeepSeek-V4-pro"):
        return "DeepSeek-V4-pro"
    if s.startswith("Qwen3.5-Plus"):
        return "Qwen3.5-Plus"
    return re.sub(r"[-_](?:Paper_)?\d+$", "", s)


def candidate_order_key(candidate: str) -> int:
    if candidate in CANDIDATE_ORDER_FORWARD:
        return CANDIDATE_ORDER_FORWARD.index(candidate)
    if candidate in FAMILY_MEAN_ORDER:
        return 100 + FAMILY_MEAN_ORDER.index(candidate)
    if candidate in FAMILY_BEST_ORDER:
        return 200 + FAMILY_BEST_ORDER.index(candidate)
    return 999


def family_order_key(family: str) -> int:
    if family in FAMILY_ORDER:
        return FAMILY_ORDER.index(family)
    return 999


# =========================================================================
# 3. File readers and case discovery
# =========================================================================


def read_text_file(path: Path) -> str:
    encodings = ["utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"]
    last_err = ""
    for enc in encodings:
        try:
            return path.read_text(encoding=enc)
        except Exception as e:
            last_err = str(e)
    raise RuntimeError(f"Failed to read text file: {path} | last_error={last_err}")


def read_docx_file(path: Path) -> str:
    if not HAS_DOCX:
        raise ImportError("python-docx is not installed. Run: pip install python-docx")
    doc = Document(str(path))
    parts = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts).strip()


def read_any_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return read_text_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    raise ValueError(f"Unsupported file type: {path}")


def is_output_file(path: Path) -> bool:
    if not path.is_file():
        return False
    if path.name.startswith("~$"):
        return False
    if path.suffix.lower() not in [".txt", ".docx"]:
        return False

    name_low = path.name.lower()
    path_str = str(path)

    if any(k.lower() in name_low for k in IGNORED_OUTPUT_KEYWORDS):
        return False
    if "_DesignJudge" in path_str or "_scores" in path_str:
        return False
    if name_low.startswith("0_run_log"):
        return False

    return True


def detect_candidate_from_filename(path: Path) -> str:
    name = path.name.lower()
    stem = path.stem.lower()
    idx = extract_index_from_filename(path)

    if "grok" in name:
        return "__IGNORE__"

    # BEAVER reports are usually docx Paper_1/2/3.
    if path.suffix.lower() == ".docx" or "beaver" in name or "paper_" in name or stem.startswith("paper"):
        return f"BEAVER_Paper_{idx}"

    if "qwen" in name:
        return f"Qwen3.5-Plus-{idx}"
    if "gemini" in name:
        return f"Gemini3.1-pro-{idx}"
    if "deepseek" in name or "deepseek" in stem:
        return f"DeepSeek-V4-pro-{idx}"
    if "gpt" in name or "5.5" in name or "gpt5" in name:
        return f"GPT5.5-{idx}"

    return path.stem


def infer_design_request(case_dir: Path) -> str:
    if case_dir.name in CASE_REQUEST_OVERRIDES:
        return CASE_REQUEST_OVERRIDES[case_dir.name]

    parts = list(case_dir.parts)
    for key, value in CASE_REQUEST_OVERRIDES.items():
        if key in parts:
            return value

    lower_path = str(case_dir).lower()
    for keywords, req in REQUEST_KEYWORDS:
        for kw in keywords:
            if kw.lower() in lower_path:
                return req

    return (
        f"Open-ended biodegradable polymer design task inferred from case folder: {case_dir.name}. "
        "Evaluate whether each candidate provides a chemically feasible, scientifically rigorous, "
        "mechanistically coherent, multi-objective design with sufficient completeness and risk awareness."
    )


def folder_has_outputs(folder: Path) -> bool:
    return any(is_output_file(p) for p in folder.iterdir()) if folder.exists() and folder.is_dir() else False


def discover_case_dirs(root_dir: Path) -> List[Path]:
    root_dir = Path(root_dir)
    case_dirs = []

    if folder_has_outputs(root_dir):
        case_dirs.append(root_dir)

    for sub in sorted([p for p in root_dir.iterdir() if p.is_dir()], key=lambda p: natural_key(p.name)):
        if sub.name.startswith("_"):
            continue
        if folder_has_outputs(sub):
            case_dirs.append(sub)

    for dg in sorted(root_dir.rglob("DesignJudge"), key=lambda p: natural_key(str(p))):
        if dg == root_dir:
            continue
        if folder_has_outputs(dg):
            case_dirs.append(dg)

    seen = set()
    uniq = []
    for d in case_dirs:
        key = str(d.resolve()).lower()
        if key not in seen:
            seen.add(key)
            uniq.append(d)
    return uniq


def collect_case_outputs(case_dir: Path) -> Dict[str, Dict[str, Any]]:
    files = sorted([p for p in case_dir.iterdir() if is_output_file(p)], key=lambda p: natural_key(p.name))
    outputs: Dict[str, Dict[str, Any]] = {}

    for fp in files:
        candidate = detect_candidate_from_filename(fp)
        if candidate == "__IGNORE__":
            continue

        try:
            txt = read_any_file(fp)
            combined = normalize_answer(f"===== SOURCE_FILE: {fp.name} =====\n\n{txt}")
        except Exception as e:
            combined = f"===== SOURCE_FILE: {fp.name} =====\n\n[READ_ERROR] {e}"

        # Avoid overwriting when duplicate names appear.
        if candidate in outputs:
            base = candidate
            k = 2
            while f"{base}_dup{k}" in outputs:
                k += 1
            candidate = f"{base}_dup{k}"

        outputs[candidate] = {
            "text": combined,
            "files": [str(fp)],
            "n_files": 1,
            "char_len": len(combined),
            "family": get_family(candidate),
        }

    return outputs


def build_cases_dataframe(root_dir: Path) -> Tuple[pd.DataFrame, Dict[str, Dict[str, str]]]:
    case_dirs = discover_case_dirs(root_dir)
    if not case_dirs:
        raise FileNotFoundError(f"No .txt/.docx outputs found under: {root_dir}")

    rows = []
    case_candidate_text: Dict[str, Dict[str, str]] = {}

    for idx, case_dir in enumerate(case_dirs, start=1):
        case_id = f"case_{idx:02d}"
        case_name = case_dir.name
        request = infer_design_request(case_dir)
        outputs = collect_case_outputs(case_dir)

        row = {
            "case_id": case_id,
            "case_name": case_name,
            "case_dir": str(case_dir),
            "design_request": request,
        }

        candidate_texts = {}
        for candidate, info in outputs.items():
            candidate_texts[candidate] = info["text"]
            row[f"{candidate}__files"] = " ; ".join(info["files"])
            row[f"{candidate}__n_files"] = info["n_files"]
            row[f"{candidate}__chars"] = info["char_len"]
            row[f"{candidate}__family"] = info["family"]
            row[candidate] = info["text"]

        case_candidate_text[case_id] = candidate_texts
        rows.append(row)

    df = pd.DataFrame(rows)
    return df, case_candidate_text


# =========================================================================
# 4. Judge prompt
# =========================================================================


def normalize_for_judge(text: str) -> str:
    if text is None:
        return ""
    s = str(text)
    s = re.sub(r"<think>.*?</think>", " ", s, flags=re.DOTALL | re.IGNORECASE)
    s = re.sub(r"```(?:json|markdown|md|text)?", " ", s, flags=re.IGNORECASE)
    s = s.replace("```", " ")

    s = re.sub(r"\$\$(.*?)\$\$", lambda m: "\n[Equation] " + m.group(1).strip() + " [/Equation]\n", s, flags=re.DOTALL)
    s = re.sub(r"\\\[(.*?)\\\]", lambda m: "\n[Equation] " + m.group(1).strip() + " [/Equation]\n", s, flags=re.DOTALL)
    s = re.sub(r"\$(.*?)\$", lambda m: " " + m.group(1).strip() + " ", s, flags=re.DOTALL)
    s = re.sub(r"\\\((.*?)\\\)", lambda m: " " + m.group(1).strip() + " ", s, flags=re.DOTALL)

    latex_map = {
        r"\\alpha": "alpha",
        r"\\beta": "beta",
        r"\\gamma": "gamma",
        r"\\delta": "delta",
        r"\\Delta": "Delta",
        r"\\mu": "mu",
        r"\\sigma": "sigma",
        r"\\approx": "approximately",
        r"\\sim": "approximately",
        r"\\times": "x",
        r"\\cdot": "*",
        r"\\pm": "+/-",
        r"\\leq": "<=",
        r"\\geq": ">=",
        r"\\to": "->",
        r"\\rightarrow": "->",
        r"\\left": "",
        r"\\right": "",
        r"\\%": "%",
    }
    for k, v in latex_map.items():
        s = re.sub(k, v, s)

    s = re.sub(r"\\text\{([^{}]*)\}", r"\1", s)
    s = re.sub(r"\\mathrm\{([^{}]*)\}", r"\1", s)
    s = re.sub(r"\\mathbf\{([^{}]*)\}", r"\1", s)

    for _ in range(5):
        new_s = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", s)
        if new_s == s:
            break
        s = new_s

    s = re.sub(r"^#{1,6}\s*", "", s, flags=re.MULTILINE)
    s = s.replace("**", "").replace("__", "").replace("`", "")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def build_candidates_block(candidate_to_prediction: Dict[str, str]) -> str:
    parts = []
    for candidate_name, prediction in candidate_to_prediction.items():
        if is_invalid_answer(prediction):
            pred = "[INVALID OR MISSING ANSWER]"
        else:
            pred = truncate_middle(normalize_for_judge(prediction), MAX_CHARS_PER_CANDIDATE)
        parts.append(f"### {candidate_name}\n{pred}")
    return "\n\n".join(parts).strip()


JUDGE_PROMPT_TEMPLATE = r"""
You are a harsh but calibrated reviewer for open-ended materials design evaluation.

Your task is to evaluate independent candidate design outputs for the same biodegradable-polymer design request.

Important context:
- This is an open-ended design task. There is NO single ground-truth design.
- Score each candidate against the original design request and the rubric below.
- Do NOT reward a candidate merely for being longer, more polished, or more report-like.
- Do NOT reward explicit literature references or citation-like formatting by itself.
- References are relevant only when they improve scientific rigor, chemical plausibility, design completeness, or risk calibration.
- Do NOT penalize raw Markdown, LaTeX-like syntax, tables, or structured formatting unless they obscure the meaning.
- Avoid position bias and model-name bias. Judge the scientific content and design quality.
- Use the full 0-10 scale. A merely plausible design is not automatically an 8+.
- Separate close candidates by 0.5-1.5 points when one is materially better.
- For the innovation/exploration dimension, consider both the novelty of the individual candidate and whether candidates from the same model family explore distinct design routes rather than minor variants of the same motif.

Original Design Request:
{design_request}

Candidate names:
{candidate_names}

Candidate Design Outputs:
{candidates_block}

IMPORTANT JSON KEY RULE:
Use exactly the candidate names listed above as the keys in the "Scores" object.
Do not rename candidates.
For example, if a candidate name is "GPT5.5-2", the JSON key must be exactly "GPT5.5-2".

Score each candidate from 0.0 to 10.0 with one decimal place on the following seven dimensions.

1) Chemical_Feasibility
Judge whether the design can realistically be made and processed in polymer/materials chemistry.
High scores require plausible monomers, polymer architecture, reaction route, composition window, processing logic, and absence of obvious synthetic impossibilities.
- 9-10: Chemically concrete, realistic, and experimentally implementable.
- 7-8.9: Mostly feasible with minor uncertainties.
- 5-6.9: Plausible but underspecified or partly questionable.
- 0-4.9: Chemically vague, unrealistic, or internally impossible.

2) Scientific_Rigor
Judge whether the reasoning is scientifically disciplined and avoids overclaiming.
High scores require correct polymer science, calibrated uncertainty, appropriate use of numbers/claims, and distinction between known facts, assumptions, and hypotheses.
- 9-10: Rigorous, well calibrated, and scientifically reliable.
- 7-8.9: Generally rigorous with minor unsupported claims.
- 5-6.9: Some correct science but notable overclaiming, weak calibration, or unsupported specificity.
- 0-4.9: Scientifically weak, misleading, or highly speculative.

3) Mechanistic_Coherence
Judge whether the structure-property-degradation mechanism is coherent.
High scores require a clear chain linking chemistry/architecture to thermal, mechanical, transport, crystallinity, hydrolysis, enzymatic degradation, or morphology outcomes.
- 9-10: Strong, coherent mechanism across structure, properties, and degradation.
- 7-8.9: Mechanistically plausible but not fully developed.
- 5-6.9: Some mechanism but shallow, incomplete, or partly disconnected.
- 0-4.9: Mostly assertion, contradiction, or no meaningful mechanism.

4) Property_Constraint_and_Tradeoff_Handling
Judge whether the design handles the target property constraints and multi-objective trade-offs.
High scores require explicit balancing of relevant constraints, such as Tg/Tm/HDT, crystallinity, toughness, modulus/strength, hydrophilicity, water uptake, degradation rate, processability, and application constraints.
- 9-10: Excellent multi-objective balancing with clear tuning knobs or composition windows.
- 7-8.9: Good trade-off handling with some concrete constraints.
- 5-6.9: Mentions trade-offs but lacks a clear balancing strategy.
- 0-4.9: Optimizes one property while ignoring major competing constraints.

5) Design_Completeness
Judge whether the output forms a complete design proposal rather than a loose idea.
High scores require candidate identity, architecture/composition, expected property changes, synthesis/fabrication route, validation plan, and optimization path.
- 9-10: Complete, experimentally actionable design package.
- 7-8.9: Mostly complete with minor missing components.
- 5-6.9: Useful proposal but missing important design or validation elements.
- 0-4.9: Fragmentary, generic, or not actionable.

6) Risk_and_Uncertainty_Awareness
Judge whether the output identifies major failure modes, uncertainty, and validation bottlenecks.
High scores require explicit discussion of risks such as slow degradation, excessive crystallinity, brittleness, phase separation, low molecular weight, processing instability, toxicity/biocompatibility, autocatalysis, or uncertain degradation kinetics, plus possible mitigation.
- 9-10: Strong risk analysis with uncertainty and mitigation.
- 7-8.9: Good risk awareness with some mitigation.
- 5-6.9: Mentions risks but incompletely.
- 0-4.9: Overly optimistic or largely ignores uncertainty.

7) Design_Innovation_and_Exploration
Judge whether the candidate proposes a nontrivial design idea and whether the corresponding model family explores a sufficiently broad design space across its three outputs.
At the candidate level, high scores require a specific and conceptually meaningful design route rather than a generic polymer modification. Examples include distinctive monomer selection, polymer architecture, morphology control, dynamic chemistry, hierarchical reinforcement, degradation-control mechanism, or multi-function coupling.
At the model-family level, also consider whether the three candidates from the same model represent genuinely distinct mechanistic routes rather than minor variations of the same motif. Outputs that repeatedly propose nearly identical hard-soft block copolymer designs should receive lower exploration credit, even if each individual candidate is plausible.
- 9-10: Highly distinctive candidate concept and broad exploration across different design mechanisms.
- 7-8.9: Clear novelty with moderate diversity among the model family's three candidates.
- 5-6.9: Some novelty, but candidates are partly generic or substantially overlapping.
- 0-4.9: Generic, repetitive, or mostly restates the target without exploring distinct design routes.

Return ONLY valid JSON in this exact schema:
{{
  "Critique": "brief overall comparative critique",
  "Scores": {{
    "CANDIDATE_NAME_1": {{
      "Chemical_Feasibility": 0.0,
      "Scientific_Rigor": 0.0,
      "Mechanistic_Coherence": 0.0,
      "Property_Constraint_and_Tradeoff_Handling": 0.0,
      "Design_Completeness": 0.0,
      "Risk_and_Uncertainty_Awareness": 0.0,
      "Design_Innovation_and_Exploration": 0.0,
      "Critique": "brief critique for this candidate"
    }},
    "CANDIDATE_NAME_2": {{
      "Chemical_Feasibility": 0.0,
      "Scientific_Rigor": 0.0,
      "Mechanistic_Coherence": 0.0,
      "Property_Constraint_and_Tradeoff_Handling": 0.0,
      "Design_Completeness": 0.0,
      "Risk_and_Uncertainty_Awareness": 0.0,
      "Design_Innovation_and_Exploration": 0.0,
      "Critique": "brief critique for this candidate"
    }}
  }}
}}
"""


def extract_json_block(text: str) -> str:
    text = clean_think_tag(text)
    block = re.search(r"```json\s*(\{.*?\})\s*```", text, flags=re.DOTALL | re.IGNORECASE)
    if block:
        return block.group(1)
    block = re.search(r"JSON\s*:\s*(\{.*?\})", text, flags=re.DOTALL | re.IGNORECASE)
    if block:
        return block.group(1)
    block = re.search(r"\{.*\}", text, flags=re.DOTALL)
    if block:
        return block.group(0)
    return ""


def calc_weighted_overall(score_like: Dict[str, Any]) -> float:
    total = 0.0
    for metric, weight in WEIGHTS.items():
        total += safe_float(score_like.get(metric, 0.0)) * weight
    return total


def normalize_score_key(key: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(key).lower())


def get_scores_for_candidate(score_map: Dict[str, Any], candidate_name: str) -> Dict[str, Any]:
    if not isinstance(score_map, dict):
        return {}

    if candidate_name in score_map and isinstance(score_map[candidate_name], dict):
        return score_map[candidate_name]

    target = normalize_score_key(candidate_name)
    for k, v in score_map.items():
        if normalize_score_key(k) == target and isinstance(v, dict):
            return v

    alias_map = {
        "BEAVER_Paper_1": ["beaver1", "beaverpaper1", "paper1", "beavercandidate1"],
        "BEAVER_Paper_2": ["beaver2", "beaverpaper2", "paper2", "beavercandidate2"],
        "BEAVER_Paper_3": ["beaver3", "beaverpaper3", "paper3", "beavercandidate3"],
    }
    for alias in alias_map.get(candidate_name, []):
        for k, v in score_map.items():
            if normalize_score_key(k) == alias and isinstance(v, dict):
                return v

    return {}


def call_openai_chat(tag: str, messages: List[Dict[str, str]], row_id: str, max_attempts: int = 3) -> Dict[str, Any]:
    if not JUDGE_CONFIG.get("api_key"):
        return {"ok": False, "content": "", "error": "Missing JUDGE_API_KEY", "seconds": 0.0}

    last_err = ""
    for attempt in range(1, max_attempts + 1):
        start = time.time()
        try:
            client = OpenAI(
                api_key=JUDGE_CONFIG["api_key"],
                base_url=JUDGE_CONFIG["base_url"],
                timeout=240.0,
                max_retries=0,
            )
            
            # 开启 stream=True 以绕过网关 60 秒超时机制
            response = client.chat.completions.create(
                model=JUDGE_CONFIG["model_name"],
                messages=messages,
                temperature=0.0,
                max_tokens=8192,
                stream=True,
                response_format={"type": "json_object"}  # 👈 新增这一行：强制 API 走 JSON 模式
            )
            
            content = ""
            reasoning_content = ""
            
            # 遍历数据流，持续接收数据块以保持连接活跃
            for chunk in response:
                if chunk.choices and len(chunk.choices) > 0:
                    delta = chunk.choices[0].delta
                    
                    # 1) 拼接标准字符串 content
                    if getattr(delta, "content", None):
                        content += delta.content
                        
                    # 2) 兼容某些推理模型 (如 DeepSeek) 可能返回的 reasoning_content
                    if getattr(delta, "reasoning_content", None):
                        reasoning_content += delta.reasoning_content
                        
                    # 3) 兼容 OpenAI SDK 严格模式下，非标准字段被塞进 model_extra 的情况
                    extra_data = getattr(delta, "model_extra", {}) or {}
                    for attr in ["reasoning", "output_text"]:
                        alt = extra_data.get(attr)
                        if isinstance(alt, str) and alt:
                            content += alt
            
            # 如果主内容为空，但推理字段有内容（极少数中转平台的特殊行为），则兜底使用推理内容
            if not content.strip() and reasoning_content.strip():
                content = reasoning_content
                
            content = normalize_answer(content)
            elapsed = time.time() - start
            
            log_msg(
                f"✅ API成功 [{tag}] case={row_id} model={JUDGE_CONFIG['model_name']} "
                f"time={elapsed:.2f}s chars={len(content)} preview={preview_text(content)}",
                also_print=False,
            )
            return {"ok": True, "content": content, "error": "", "seconds": elapsed}
            
        except Exception as e:
            elapsed = time.time() - start
            last_err = str(e)
            log_msg(
                f"❌ API失败 [{tag}] case={row_id} attempt={attempt}/{max_attempts} "
                f"time={elapsed:.2f}s error={last_err}",
                also_print=False,
            )
            if attempt < max_attempts:
                time.sleep(min(20, 3 * attempt))

    return {"ok": False, "content": "", "error": last_err, "seconds": 0.0}
    
def build_record(
    case_id: str,
    case_name: str,
    case_dir: str,
    design_request: str,
    order_run: str,
    candidate_name: str,
    payload: Dict[str, Any],
) -> Dict[str, Any]:
    record = {
        "case_id": case_id,
        "case_name": case_name,
        "case_dir": case_dir,
        "design_request": design_request,
        "OrderRun": order_run,
        "Candidate": candidate_name,
        "Family": get_family(candidate_name),
        "JudgeParseOK": int(payload.get("JudgeParseOK", 0)),
        "Critique": payload.get("Critique", ""),
    }
    for metric in METRIC_COLUMNS:
        record[metric] = round(max(0.0, min(10.0, safe_float(payload.get(metric, 0.0)))), 1)
    record["Overall"] = round(calc_weighted_overall(record), 3)
    return record


def evaluate_case_all_candidates(
    case_id: str,
    case_name: str,
    case_dir: str,
    design_request: str,
    candidate_to_prediction: Dict[str, str],
    order_run: str,
) -> Tuple[str, str, List[Dict[str, Any]], bool]:

    default_payload = {metric: 0.0 for metric in METRIC_COLUMNS}
    default_payload.update({"JudgeParseOK": 0, "Critique": "Judge failed or answer invalid."})

    candidate_to_prediction_for_judge = {}
    for candidate_name, prediction in candidate_to_prediction.items():
        candidate_to_prediction_for_judge[candidate_name] = prediction if not is_invalid_answer(prediction) else "[INVALID OR MISSING ANSWER]"

    candidate_names = "\n".join([f"- {m}" for m in candidate_to_prediction_for_judge.keys()])
    prompt = JUDGE_PROMPT_TEMPLATE.format(
        design_request=normalize_for_judge(design_request),
        candidate_names=candidate_names,
        candidates_block=build_candidates_block(candidate_to_prediction_for_judge),
    )

    # 动态构建 message，不污染你原本的 prompt 模板
    system_instruction = (
        "CRITICAL: You are an automated JSON API. "
        "DO NOT output any reasoning, thinking process, preamble, or conversational text. "
        "YOUR ENTIRE RESPONSE MUST BE ONLY VALID JSON."
    )

    result = call_openai_chat(
        tag=f"DesignJudge::{order_run}",
        messages=[
            {"role": "system", "content": system_instruction},  # 👈 动态注入系统指令闭它的嘴
            {"role": "user", "content": prompt}
        ],
        row_id=case_id,
        max_attempts=3,
    )

    records = []

    if not result["ok"]:
        for candidate_name in candidate_to_prediction.keys():
            payload = default_payload.copy()
            payload["Critique"] = result.get("error", "Judge API failed.")
            records.append(build_record(case_id, case_name, case_dir, design_request, order_run, candidate_name, payload))
        return case_id, order_run, records, False

    json_str = extract_json_block(result["content"])

    if not json_str:
        log_msg(f"❌ Judge解析失败 case={case_id} run={order_run} raw={preview_text(result['content'])}", also_print=False)
        for candidate_name in candidate_to_prediction.keys():
            payload = default_payload.copy()
            payload["Critique"] = "Judge returned no valid JSON."
            records.append(build_record(case_id, case_name, case_dir, design_request, order_run, candidate_name, payload))
        return case_id, order_run, records, False

    try:
        parsed = json.loads(json_str)
        score_map = parsed.get("Scores", {}) if isinstance(parsed, dict) else {}

        for candidate_name, prediction in candidate_to_prediction.items():
            if is_invalid_answer(prediction):
                payload = default_payload.copy()
                payload["Critique"] = "Invalid or missing candidate output."
                records.append(build_record(case_id, case_name, case_dir, design_request, order_run, candidate_name, payload))
                continue

            scores = get_scores_for_candidate(score_map, candidate_name)
            missing_metrics = [m for m in METRIC_COLUMNS if m not in scores]

            if not scores or missing_metrics:
                payload = default_payload.copy()
                payload["Critique"] = (
                    f"Judge returned JSON but did not provide complete scores for candidate '{candidate_name}'. "
                    f"Missing metrics: {missing_metrics}. Available keys: {list(score_map.keys()) if isinstance(score_map, dict) else 'N/A'}"
                )
                records.append(build_record(case_id, case_name, case_dir, design_request, order_run, candidate_name, payload))
                continue

            payload = {}
            for metric in METRIC_COLUMNS:
                payload[metric] = round(max(0.0, min(10.0, safe_float(scores.get(metric, 0.0)))), 1)
            payload["Critique"] = str(scores.get("Critique", parsed.get("Critique", ""))).strip()
            payload["JudgeParseOK"] = 1
            records.append(build_record(case_id, case_name, case_dir, design_request, order_run, candidate_name, payload))

        return case_id, order_run, records, True

    except Exception as e:
        log_msg(f"❌ Judge JSON解析异常 case={case_id} run={order_run} error={e} raw={preview_text(result['content'])}", also_print=False)
        for candidate_name in candidate_to_prediction.keys():
            payload = default_payload.copy()
            payload["Critique"] = f"Judge JSON parse failed: {e}"
            records.append(build_record(case_id, case_name, case_dir, design_request, order_run, candidate_name, payload))
        return case_id, order_run, records, False


# =========================================================================
# 5. Summary and plotting
# =========================================================================


def average_forward_reverse(df_forward: pd.DataFrame, df_reverse: pd.DataFrame) -> pd.DataFrame:
    df_all = pd.concat([df_forward, df_reverse], ignore_index=True)
    if df_all.empty:
        return pd.DataFrame(columns=ITEM_SCORE_COLUMNS)

    key_cols = ["case_id", "case_name", "case_dir", "design_request", "Candidate", "Family"]
    agg = {metric: "mean" for metric in METRIC_COLUMNS}
    agg["Overall"] = "mean"
    agg["JudgeParseOK"] = "mean"
    agg["Critique"] = lambda x: " || ".join([str(v) for v in x if str(v).strip()])[:2000]

    df_avg = df_all.groupby(key_cols, dropna=False).agg(agg).reset_index()
    df_avg["OrderRun"] = "average_forward_reverse"

    for metric in METRIC_COLUMNS:
        df_avg[metric] = df_avg[metric].round(3)
    df_avg["Overall"] = df_avg["Overall"].round(3)
    df_avg["JudgeParseOK"] = (df_avg["JudgeParseOK"] * 100).round(1)

    df_avg = df_avg[ITEM_SCORE_COLUMNS]
    return df_avg


def build_final_by_candidate(df_avg: pd.DataFrame) -> pd.DataFrame:
    if df_avg.empty:
        return pd.DataFrame()

    rows = []
    for candidate, sub in df_avg.groupby("Candidate"):
        row = {
            "Candidate": candidate,
            "Family": get_family(candidate),
            "Overall": round(sub["Overall"].mean(), 3),
            "ScoredCases": int(len(sub)),
            "JudgeParseOK(%)": round(sub["JudgeParseOK"].mean(), 1),
        }
        for metric in METRIC_COLUMNS:
            row[metric] = round(sub[metric].mean(), 3)
        rows.append(row)

    df = pd.DataFrame(rows)
    if not df.empty:
        df["__order"] = df["Candidate"].map(candidate_order_key)
        df = df.sort_values(["__order", "Candidate"]).drop(columns="__order").reset_index(drop=True)
    return df


def build_family_rows(df_avg: pd.DataFrame, mode: str = "mean") -> pd.DataFrame:
    """
    mode='mean': family score = mean of all candidate scores in that family per case.
    mode='best': family score = best candidate score in that family per case.
    """
    if df_avg.empty:
        return pd.DataFrame(columns=ITEM_SCORE_COLUMNS)

    rows = []
    for keys, sub in df_avg.groupby(["case_id", "case_name", "case_dir", "design_request", "Family"], dropna=False):
        case_id, case_name, case_dir, design_request, family = keys
        sub = sub.copy()
        if family not in FAMILY_ORDER:
            continue

        row = {
            "case_id": case_id,
            "case_name": case_name,
            "case_dir": case_dir,
            "design_request": design_request,
            "OrderRun": f"family_{mode}_from_candidate_average",
            "Candidate": f"{family}_{'Mean' if mode == 'mean' else 'Best'}",
            "Family": family,
        }

        if mode == "best":
            best = sub.sort_values("Overall", ascending=False).iloc[0]
            for metric in METRIC_COLUMNS:
                row[metric] = round(float(best[metric]), 3)
            row["Overall"] = round(float(best["Overall"]), 3)
            row["JudgeParseOK"] = round(float(best["JudgeParseOK"]), 1)
            row["Critique"] = f"Derived best candidate for {family}: {best['Candidate']}"
        else:
            for metric in METRIC_COLUMNS:
                row[metric] = round(sub[metric].mean(), 3)
            row["Overall"] = round(sub["Overall"].mean(), 3)
            row["JudgeParseOK"] = round(sub["JudgeParseOK"].mean(), 1)
            row["Critique"] = f"Derived mean of {len(sub)} candidates for {family}."

        rows.append(row)

    out = pd.DataFrame(rows)
    if out.empty:
        return pd.DataFrame(columns=ITEM_SCORE_COLUMNS)
    return out[ITEM_SCORE_COLUMNS]


def build_final_by_family(family_rows: pd.DataFrame, suffix: str) -> pd.DataFrame:
    if family_rows.empty:
        return pd.DataFrame()

    rows = []
    for family, sub in family_rows.groupby("Family"):
        row = {
            "Family": family,
            "Candidate": f"{family}_{suffix}",
            "Overall": round(sub["Overall"].mean(), 3),
            "ScoredCases": int(len(sub)),
            "JudgeParseOK(%)": round(sub["JudgeParseOK"].mean(), 1),
        }
        for metric in METRIC_COLUMNS:
            row[metric] = round(sub[metric].mean(), 3)
        rows.append(row)

    df = pd.DataFrame(rows)
    if not df.empty:
        df["__order"] = df["Family"].map(family_order_key)
        df = df.sort_values("__order").drop(columns="__order").reset_index(drop=True)
    return df


def build_overall_matrix(df: pd.DataFrame, column: str = "Candidate") -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    mat = df.pivot(index="case_name", columns=column, values="Overall").reset_index()
    cols = [c for c in mat.columns if c != "case_name"]
    if column == "Candidate":
        order = [c for c in CANDIDATE_ORDER_FORWARD + FAMILY_MEAN_ORDER + FAMILY_BEST_ORDER if c in cols]
    else:
        order = [f for f in FAMILY_ORDER if f in cols]
    ordered = order + [c for c in cols if c not in order]
    return mat[["case_name"] + ordered]


def plot_metric_heatmap(df_scores: pd.DataFrame, index_col: str, save_path: Path, title: str):
    if not HAS_MPL or df_scores.empty:
        return

    data = df_scores.set_index(index_col)[METRIC_COLUMNS]
    plt.figure(figsize=(14, max(5, 0.45 * len(data))))

    if HAS_SEABORN:
        sns.heatmap(
            data,
            annot=True,
            fmt=".2f",
            vmin=0,
            vmax=10,
            linewidths=0.5,
            linecolor="white",
            cbar_kws={"label": "Score (0-10)"},
        )
    else:
        ax = plt.gca()
        im = ax.imshow(data.values, aspect="auto", vmin=0, vmax=10)
        plt.colorbar(im, ax=ax, label="Score (0-10)")
        ax.set_xticks(np.arange(len(data.columns)))
        ax.set_xticklabels(data.columns, rotation=35, ha="right")
        ax.set_yticks(np.arange(len(data.index)))
        ax.set_yticklabels(data.index)
        for i in range(data.shape[0]):
            for j in range(data.shape[1]):
                ax.text(j, i, f"{data.iloc[i, j]:.2f}", ha="center", va="center", fontsize=8)

    plt.title(title, fontweight="bold")
    plt.xlabel("")
    plt.ylabel("")
    plt.xticks(rotation=35, ha="right")
    plt.tight_layout()
    plt.savefig(save_path, dpi=600)
    plt.close()
    log_msg(f"📊 heatmap saved: {save_path}", also_print=False)


def plot_overall_bar(df_scores: pd.DataFrame, label_col: str, save_path: Path, title: str):
    if not HAS_MPL or df_scores.empty:
        return

    df = df_scores.copy()
    if label_col == "Candidate":
        df["__order"] = df[label_col].map(candidate_order_key)
    elif label_col == "Family":
        df["__order"] = df[label_col].map(family_order_key)
    else:
        df["__order"] = range(len(df))
    df = df.sort_values(["__order", label_col]).drop(columns="__order")

    x = np.arange(len(df))
    plt.figure(figsize=(max(10, 0.7 * len(df)), 5.8))
    plt.bar(x, df["Overall"].values, edgecolor="black", linewidth=1.0)

    for i, val in enumerate(df["Overall"].values):
        plt.text(i, val + 0.08, f"{val:.2f}", ha="center", va="bottom", fontsize=8)

    plt.xticks(x, df[label_col].values, rotation=30, ha="right")
    plt.ylabel("Weighted overall score (0-10)")
    plt.ylim(0, 10.8)
    plt.title(title, fontweight="bold")
    plt.tight_layout()
    plt.savefig(save_path, dpi=600)
    plt.close()
    log_msg(f"📊 overall plot saved: {save_path}", also_print=False)


# =========================================================================
# 5b. Intra-family similarity and exploration analysis
# =========================================================================


def _simple_tokenize_for_similarity(text: str) -> List[str]:
    """Lightweight tokenizer for fallback similarity calculation."""
    s = normalize_for_judge(text or "").lower()
    tokens = re.findall(r"[a-z][a-z0-9_\-]{2,}|[\u4e00-\u9fff]", s)
    stop = {
        "the", "and", "for", "with", "that", "this", "from", "into", "are", "can",
        "will", "should", "polymer", "design", "candidate", "material", "property",
        "properties", "synthesis", "route", "table", "figure", "abstract", "conclusion",
    }
    return [t for t in tokens if t not in stop]


def _fallback_cosine_similarity(texts: List[str]) -> np.ndarray:
    """Compute cosine similarity using simple term-frequency vectors, without sklearn."""
    token_lists = [_simple_tokenize_for_similarity(t) for t in texts]
    vocab = {}
    for toks in token_lists:
        for tok in toks:
            if tok not in vocab:
                vocab[tok] = len(vocab)
    if not vocab:
        return np.eye(len(texts))

    mat = np.zeros((len(texts), len(vocab)), dtype=float)
    for i, toks in enumerate(token_lists):
        for tok in toks:
            mat[i, vocab[tok]] += 1.0
    norms = np.linalg.norm(mat, axis=1)
    norms[norms == 0] = 1.0
    mat = mat / norms[:, None]
    return mat @ mat.T


def compute_text_similarity_matrix(texts: List[str]) -> np.ndarray:
    """Compute semantic-ish text similarity. Uses sklearn TF-IDF if available, otherwise fallback TF cosine."""
    if len(texts) == 0:
        return np.zeros((0, 0), dtype=float)
    if len(texts) == 1:
        return np.ones((1, 1), dtype=float)

    cleaned = [normalize_for_judge(t or "") for t in texts]
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity

        vectorizer = TfidfVectorizer(
            lowercase=True,
            ngram_range=(1, 2),
            min_df=1,
            max_df=0.95,
            stop_words="english",
            token_pattern=r"(?u)\b[a-zA-Z][a-zA-Z0-9_\-]{2,}\b",
        )
        x = vectorizer.fit_transform(cleaned)
        return cosine_similarity(x)
    except Exception:
        return _fallback_cosine_similarity(cleaned)


def build_family_exploration_table(
    df_avg: pd.DataFrame,
    case_candidate_text: Dict[str, Dict[str, str]],
) -> pd.DataFrame:
    """
    Compute intra-family candidate similarity and design-space exploration.

    MeanSimilarity: average pairwise text similarity among the candidates of the same family.
    DiversityRaw: 1 - MeanSimilarity.
    Design_Space_Exploration_Score: min-max normalized DiversityRaw to 0-10 within each case.
    """
    rows = []
    if df_avg.empty:
        return pd.DataFrame()

    for keys, sub in df_avg.groupby(["case_id", "case_name", "case_dir", "design_request", "Family"], dropna=False):
        case_id, case_name, case_dir, design_request, family = keys
        if family not in FAMILY_ORDER:
            continue

        candidates = [c for c in sub["Candidate"].astype(str).tolist() if not c.endswith("_Mean") and not c.endswith("_Best")]
        texts = []
        valid_candidates = []
        for cand in candidates:
            txt = case_candidate_text.get(str(case_id), {}).get(cand, "")
            if not is_invalid_answer(txt):
                valid_candidates.append(cand)
                texts.append(txt)

        if len(texts) >= 2:
            sim = compute_text_similarity_matrix(texts)
            pair_vals = []
            pair_labels = []
            for i in range(len(texts)):
                for j in range(i + 1, len(texts)):
                    pair_vals.append(float(sim[i, j]))
                    pair_labels.append(f"{valid_candidates[i]} vs {valid_candidates[j]}={sim[i, j]:.4f}")
            mean_sim = float(np.mean(pair_vals)) if pair_vals else 1.0
        elif len(texts) == 1:
            mean_sim = 1.0
            pair_labels = []
        else:
            mean_sim = np.nan
            pair_labels = []

        diversity_raw = 1.0 - mean_sim if not pd.isna(mean_sim) else np.nan
        rows.append({
            "case_id": case_id,
            "case_name": case_name,
            "case_dir": case_dir,
            "design_request": design_request,
            "Family": family,
            "N_Candidates": len(texts),
            "MeanPairwiseSimilarity": round(mean_sim, 4) if not pd.isna(mean_sim) else np.nan,
            "DiversityRaw_1_minus_similarity": round(diversity_raw, 4) if not pd.isna(diversity_raw) else np.nan,
            "PairwiseDetails": " ; ".join(pair_labels),
        })

    df = pd.DataFrame(rows)
    if df.empty:
        return df

    df["Design_Space_Exploration_Score"] = np.nan
    for case_id, sub_idx in df.groupby("case_id").groups.items():
        vals = df.loc[sub_idx, "DiversityRaw_1_minus_similarity"].astype(float)
        vmin = vals.min(skipna=True)
        vmax = vals.max(skipna=True)
        if pd.isna(vmin) or pd.isna(vmax):
            continue
        if abs(vmax - vmin) < 1e-12:
            df.loc[sub_idx, "Design_Space_Exploration_Score"] = 5.0
        else:
            df.loc[sub_idx, "Design_Space_Exploration_Score"] = 10.0 * (vals - vmin) / (vmax - vmin)

    df["Design_Space_Exploration_Score"] = df["Design_Space_Exploration_Score"].round(3)
    return df


def merge_exploration_into_family_mean(
    df_family_mean: pd.DataFrame,
    df_exploration: pd.DataFrame,
    exploration_weight: float = 0.10,
) -> pd.DataFrame:
    """Add objective similarity-derived exploration score and an optional adjusted overall."""
    if df_family_mean.empty or df_exploration.empty:
        return df_family_mean.copy()

    exp_summary = df_exploration.groupby("Family", as_index=False).agg({
        "MeanPairwiseSimilarity": "mean",
        "DiversityRaw_1_minus_similarity": "mean",
        "Design_Space_Exploration_Score": "mean",
    })
    exp_summary = exp_summary.rename(columns={
        "MeanPairwiseSimilarity": "MeanPairwiseSimilarity_AcrossCases",
        "DiversityRaw_1_minus_similarity": "DiversityRaw_AcrossCases",
        "Design_Space_Exploration_Score": "Objective_Design_Space_Exploration",
    })
    for c in ["MeanPairwiseSimilarity_AcrossCases", "DiversityRaw_AcrossCases", "Objective_Design_Space_Exploration"]:
        exp_summary[c] = exp_summary[c].round(3)

    out = df_family_mean.merge(exp_summary, on="Family", how="left")
    out["Exploration_Adjusted_Overall"] = (
        (1.0 - exploration_weight) * out["Overall"].astype(float)
        + exploration_weight * out["Objective_Design_Space_Exploration"].fillna(0).astype(float)
    ).round(3)
    out["Exploration_Adjustment_Weight"] = exploration_weight
    return out


# =========================================================================
# 6. Main scoring
# =========================================================================


def build_ordered_candidate_texts(candidate_texts: Dict[str, str], order: List[str]) -> Dict[str, str]:
    ordered = {}
    for candidate in order:
        if candidate in candidate_texts and not is_invalid_answer(candidate_texts[candidate]):
            ordered[candidate] = candidate_texts[candidate]

    for candidate, txt in candidate_texts.items():
        if candidate not in ordered and not is_invalid_answer(txt):
            ordered[candidate] = txt
    return ordered


def run_order_scoring(
    df_cases: pd.DataFrame,
    case_candidate_text: Dict[str, Dict[str, str]],
    order_run: str,
    order: List[str],
) -> pd.DataFrame:

    out_path = FORWARD_SCORE_PATH if order_run == "forward" else REVERSE_SCORE_PATH

    if out_path.exists() and not FORCE_REJUDGE:
        try:
            old = pd.read_csv(out_path)
            expected_pairs = set()
            for cid in df_cases["case_id"].astype(str):
                ordered_candidates = build_ordered_candidate_texts(case_candidate_text.get(cid, {}), order)
                for c in ordered_candidates.keys():
                    expected_pairs.add((cid, c))
            old_pairs = set(zip(old["case_id"].astype(str), old["Candidate"].astype(str)))
            if expected_pairs.issubset(old_pairs):
                log_msg(f"⏩ {order_run} scores already complete, loaded: {out_path}")
                return old
        except Exception as e:
            log_msg(f"⚠️ Failed to load existing {order_run} scores; will re-score. error={e}")

    records_all: List[Dict[str, Any]] = []
    futures = []

    with concurrent.futures.ThreadPoolExecutor(max_workers=MAX_WORKERS_JUDGE) as executor:
        for _, row in df_cases.iterrows():
            case_id = str(row["case_id"])
            candidate_texts = build_ordered_candidate_texts(case_candidate_text.get(case_id, {}), order)

            if len(candidate_texts) < 2:
                log_msg(f"⚠️ Skipping {case_id}: fewer than 2 valid candidate outputs found. candidates={list(candidate_texts.keys())}")
                continue

            futures.append(
                executor.submit(
                    evaluate_case_all_candidates,
                    case_id,
                    str(row["case_name"]),
                    str(row["case_dir"]),
                    str(row["design_request"]),
                    candidate_texts,
                    order_run,
                )
            )

        done_count = 0
        for future in tqdm(concurrent.futures.as_completed(futures), total=len(futures), desc=f"Judge::{order_run}"):
            try:
                case_id, run, recs, ok = future.result()
                records_all.extend(recs)
                done_count += 1

                if done_count % SAVE_EVERY_N_JUDGE == 0 or done_count == len(futures):
                    df_tmp = pd.DataFrame(records_all)
                    if not df_tmp.empty:
                        df_tmp = df_tmp[ITEM_SCORE_COLUMNS]
                        save_df_atomic(df_tmp, out_path)
                    log_msg(f"💾 Saved {order_run} scores: {done_count}/{len(futures)} latest_ok={ok}")
            except Exception as e:
                log_msg(f"❌ Future failed in {order_run}: {e}\n{traceback.format_exc()}")

    df_scores = pd.DataFrame(records_all)
    if not df_scores.empty:
        df_scores = df_scores[ITEM_SCORE_COLUMNS]
        save_df_atomic(df_scores, out_path)
    return df_scores


def main():
    log_msg("=" * 80)
    log_msg("Start Design-mode 3-candidates-per-family LLM-as-a-Judge scoring")
    log_msg(f"ROOT_DIR = {ROOT_DIR}")
    log_msg(f"OUTPUT_DIR = {OUTPUT_DIR}")
    log_msg(f"Judge model = {JUDGE_CONFIG.get('model_name')}")
    log_msg(f"Judge base_url = {JUDGE_CONFIG.get('base_url')}")
    log_msg(f"Judge api_key_present = {bool(JUDGE_CONFIG.get('api_key'))}")
    log_msg(f"MAX_CHARS_PER_CANDIDATE = {MAX_CHARS_PER_CANDIDATE}")
    log_msg(f"Ignored output keywords = {IGNORED_OUTPUT_KEYWORDS}")
    log_msg(f"Metrics = {METRIC_COLUMNS}")
    log_msg(f"Weights = {WEIGHTS}")

    if not ROOT_DIR.exists():
        raise FileNotFoundError(f"ROOT_DIR does not exist: {ROOT_DIR}")

    if not JUDGE_CONFIG.get("api_key"):
        raise RuntimeError(
            "Missing JUDGE_API_KEY. Set it before running, for example in PowerShell:\n"
            '$env:JUDGE_API_KEY="your_api_key"\n'
            '$env:JUDGE_BASE_URL="your_api_url"\n'
            '$env:JUDGE_MODEL="deepseek-v4-pro"\n'
            "Then run: python Design_BEAVER_vs_Models_deepseek.py"
        )

    df_cases, case_candidate_text = build_cases_dataframe(ROOT_DIR)

    text_candidate_cols = set()
    for _, cand_map in case_candidate_text.items():
        text_candidate_cols.update(cand_map.keys())

    df_discovered = df_cases.drop(columns=[c for c in df_cases.columns if c in text_candidate_cols], errors="ignore")
    save_df_atomic(df_discovered, DISCOVERED_CASES_PATH)
    save_df_atomic(df_cases, RAW_OUTPUTS_PATH)

    log_msg("=== Discovered cases ===")
    for _, row in df_cases.iterrows():
        cid = str(row["case_id"])
        candidates_present = list(case_candidate_text.get(cid, {}).keys())
        ordered_present = [c for c in CANDIDATE_ORDER_FORWARD if c in candidates_present]
        extra_present = [c for c in candidates_present if c not in CANDIDATE_ORDER_FORWARD]
        log_msg(
            f"{cid} | {row['case_name']} | n_candidates={len(candidates_present)} | "
            f"ordered={ordered_present} | extra={extra_present} | request={row['design_request']}"
        )

    df_forward = run_order_scoring(df_cases, case_candidate_text, "forward", CANDIDATE_ORDER_FORWARD)
    df_reverse = run_order_scoring(df_cases, case_candidate_text, "reverse", CANDIDATE_ORDER_REVERSE)

    df_avg = average_forward_reverse(df_forward, df_reverse)
    save_df_atomic(df_avg, AVERAGE_SCORE_PATH)

    df_by_candidate = build_final_by_candidate(df_avg)
    save_df_atomic(df_by_candidate, FINAL_BY_CANDIDATE_PATH)

    family_mean_rows = build_family_rows(df_avg, mode="mean")
    family_best_rows = build_family_rows(df_avg, mode="best")

    df_family_mean = build_final_by_family(family_mean_rows, suffix="Mean")
    df_family_best = build_final_by_family(family_best_rows, suffix="Best")
    save_df_atomic(df_family_mean, FINAL_BY_FAMILY_MEAN_PATH)
    save_df_atomic(df_family_best, FINAL_BY_FAMILY_BEST_PATH)

    df_exploration = build_family_exploration_table(df_avg, case_candidate_text)
    save_df_atomic(df_exploration, EXPLORATION_PATH)

    df_family_mean_with_exploration = merge_exploration_into_family_mean(
        df_family_mean,
        df_exploration,
        exploration_weight=0.10,
    )
    save_df_atomic(df_family_mean_with_exploration, FAMILY_MEAN_WITH_EXPLORATION_PATH)

    df_matrix_candidate = build_overall_matrix(df_avg, column="Candidate")
    save_df_atomic(df_matrix_candidate, OVERALL_MATRIX_CANDIDATE_PATH)

    df_matrix_family_mean = build_overall_matrix(family_mean_rows, column="Candidate")
    save_df_atomic(df_matrix_family_mean, OVERALL_MATRIX_FAMILY_MEAN_PATH)

    critique_cols = ["case_id", "case_name", "Candidate", "Family", "OrderRun", "Critique"]
    df_critiques = pd.concat([df_forward, df_reverse], ignore_index=True)
    if not df_critiques.empty:
        save_df_atomic(df_critiques[critique_cols], CRITIQUE_PATH)

    plot_metric_heatmap(
        df_by_candidate,
        index_col="Candidate",
        save_path=HEATMAP_CANDIDATE_PATH,
        title="Design-mode seven-dimensional score heatmap by candidate",
    )
    plot_metric_heatmap(
        df_family_mean,
        index_col="Family",
        save_path=HEATMAP_FAMILY_MEAN_PATH,
        title="Design-mode seven-dimensional score heatmap by family mean",
    )
    plot_overall_bar(
        df_by_candidate,
        label_col="Candidate",
        save_path=OVERALL_CANDIDATE_PLOT_PATH,
        title="Design-mode overall score by candidate",
    )
    plot_overall_bar(
        df_family_mean,
        label_col="Family",
        save_path=OVERALL_FAMILY_MEAN_PLOT_PATH,
        title="Design-mode overall score by family mean",
    )

    print("\n" + "=" * 80)
    print("🏆 Final scores by candidate")
    print("=" * 80)
    print(df_by_candidate.to_string(index=False) if not df_by_candidate.empty else "No available scores.")

    print("\n" + "=" * 80)
    print("🏆 Final scores by family mean")
    print("=" * 80)
    print(df_family_mean.to_string(index=False) if not df_family_mean.empty else "No available family scores.")

    print("\n" + "=" * 80)
    print("🏆 Final scores by family mean with objective exploration")
    print("=" * 80)
    print(df_family_mean_with_exploration.to_string(index=False) if not df_family_mean_with_exploration.empty else "No available family exploration scores.")

    print("\n" + "=" * 80)
    print("🏆 Final scores by family best")
    print("=" * 80)
    print(df_family_best.to_string(index=False) if not df_family_best.empty else "No available family best scores.")

    print("\n" + "=" * 80)
    print("📌 Overall matrix by candidate")
    print("=" * 80)
    print(df_matrix_candidate.to_string(index=False) if not df_matrix_candidate.empty else "No matrix.")

    log_msg(f"✅ discovered cases saved: {DISCOVERED_CASES_PATH}")
    log_msg(f"✅ raw outputs saved: {RAW_OUTPUTS_PATH}")
    log_msg(f"✅ forward scores saved: {FORWARD_SCORE_PATH}")
    log_msg(f"✅ reverse scores saved: {REVERSE_SCORE_PATH}")
    log_msg(f"✅ average per-candidate scores saved: {AVERAGE_SCORE_PATH}")
    log_msg(f"✅ final by candidate saved: {FINAL_BY_CANDIDATE_PATH}")
    log_msg(f"✅ final by family mean saved: {FINAL_BY_FAMILY_MEAN_PATH}")
    log_msg(f"✅ final by family best saved: {FINAL_BY_FAMILY_BEST_PATH}")
    log_msg(f"✅ family similarity/exploration saved: {EXPLORATION_PATH}")
    log_msg(f"✅ family mean with objective exploration saved: {FAMILY_MEAN_WITH_EXPLORATION_PATH}")
    log_msg(f"✅ overall candidate matrix saved: {OVERALL_MATRIX_CANDIDATE_PATH}")
    log_msg(f"✅ overall family mean matrix saved: {OVERALL_MATRIX_FAMILY_MEAN_PATH}")
    if HAS_MPL:
        log_msg(f"✅ candidate heatmap saved: {HEATMAP_CANDIDATE_PATH}")
        log_msg(f"✅ family mean heatmap saved: {HEATMAP_FAMILY_MEAN_PATH}")
        log_msg(f"✅ candidate overall plot saved: {OVERALL_CANDIDATE_PLOT_PATH}")
        log_msg(f"✅ family mean overall plot saved: {OVERALL_FAMILY_MEAN_PLOT_PATH}")
    log_msg("Finished")
    log_msg("=" * 80)


if __name__ == "__main__":
    main()
